from django.shortcuts import render, render_to_response
from django.http import HttpResponseRedirect, HttpResponse, Http404
from django.utils import timezone
from django.template.context_processors import csrf
from forms import FileForm
from models import CSVFile, germline_snp, somatic_snp
from django.contrib.auth.decorators import login_required
import vcf
import re
import requests
import sys
import csv
import json
from requests.exceptions import ConnectionError
from django.core.files import File
from docx import *
from docx.shared import Inches
from django.conf import settings
import datetime
from pp.client.python.unoconv import unoconv
import subprocess

# function to upload VCF file and then to generate annotated SNP file as a result
def file_form(request):
    # if not request.user.is_staff or not request.user.is_superuser:
    #     raise Http404;
    if request.POST:
        form = FileForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.user = request.user
            instance.save()

            # read CSV file and define parameters
            input_file_name = re.search(r'/(media/.*)', instance.file.url)
            fileID = instance.id
            input_file = vcf.Reader(open(input_file_name.group(1), "r"))
            # output_somatic_file = open('media/somatic_mutation-'+str(fileID)+'.tsv', 'wb')
            # output_germline_file = open('media/germline_mutation-'+str(fileID)+'.tsv', 'wb')
            logs = open('media/logs-'+str(fileID)+'.txt', 'w')
            somatic_mutation_array = []
            germline_mutation_array = []

            # Read vcf file and create seperate arrays for somatic and germline mutation
            for record in input_file:
                chr = re.search(r'chr(.*)', str(record.CHROM))
                start_pos = str(record.POS)
                if record.REF == '-':
                    end_pos = record.POS + 1
                else:
                    end_pos = start_pos
                ref = str(record.REF)
                alt = re.search(r'\[(.*)\]', str(record.ALT))
                oncoTator = 'http://www.broadinstitute.org/oncotator/mutation/' + chr.group(
                    1) + '_' + start_pos + '_' + end_pos + '_' + ref + '_' + alt.group(1)
                try:
                    oncoTator_response = requests.get(oncoTator)
                except ConnectionError as e:
                    print e
                    logs.write(str(e))
                    oncoTator_response = "No response"
                if oncoTator_response != "No response" and oncoTator_response.status_code != 500:
                    mutation = json.loads(oncoTator_response.text)
                    if "ERROR" in mutation:
                        logs.write("OncoTator is giving error for url:" + oncoTator + " Error message: " + str(oncoTator_response.text))
                    else:

                        if mutation["CGC_Cancer Germline Mut"] == "yes":
                            germline_mutation_array.append(mutation)

                        if mutation["CGC_Cancer Somatic Mut"] == "yes":
                            somatic_mutation_array.append(mutation)
            logs.close()

            # write values from arrays to tsv files
            # fieldnames = []
            # for key in somatic_mutation_array[0]:
            #     fieldnames.append(str(key))

            # tsv_somatic_file = csv.DictWriter(output_somatic_file, fieldnames=fieldnames, delimiter='\t')
            # tsv_somatic_file.writeheader()
            for mutation in somatic_mutation_array:
                # tsv_somatic_file.writerow(mutation)
                ssnp = somatic_snp.objects.create(study_name=instance, gene=mutation["gene"], chromosome=mutation["ESP_Chromosome"], start=mutation["start"], end=mutation["end"], ref_allele=mutation["ref_allele"], alt_allele=mutation["alt_allele"], variant_class=mutation["variant_classification"],variant_type=mutation["variant_type"], ensemblID=mutation["HGNC_Ensembl Gene ID"])
            # output_somatic_file.close()

            # tsv_germline_file = csv.DictWriter(output_germline_file, fieldnames=fieldnames, delimiter='\t')
            # tsv_germline_file.writeheader()
            for mutation in germline_mutation_array:
                # tsv_germline_file.writerow(mutation)
                gsnp = germline_snp.objects.create(study_name=instance, gene=mutation["gene"], chromosome=mutation["ESP_Chromosome"], start=mutation["start"], end=mutation["end"], ref_allele=mutation["ref_allele"], alt_allele=mutation["alt_allele"], variant_class=mutation["variant_classification"],variant_type=mutation["variant_type"], ensemblID=mutation["HGNC_Ensembl Gene ID"])
            # output_germline_file.close()

            # fileread = open('media/germline_mutation-'+str(fileID)+'.tsv', 'r')
            # TSVfile = File(fileread)
            # instance.germline_result.save('media/germline_mutation-'+str(fileID)+'.tsv', TSVfile, save=True)
            #
            # fileread = open('media/somatic_mutation-' + str(fileID) + '.tsv', 'r')
            # TSVfile = File(fileread)
            # instance.somatic_result.save('media/somtic_mutation-' + str(fileID) + '.tsv', TSVfile, save=True)

            return HttpResponseRedirect('/result/')
    else:
        form = FileForm()
    return render(request, 'oncotator/snp_tools.html', {'form': form})

# result page with all studies and results, first page after login
@login_required(login_url='/login/')
def result(request):
    files = CSVFile.objects.all()
    return render(request, 'oncotator/result.html', {'files': files})

# Somatic SNP detail page, displayed when clicked on results

def somatic_snp_detail(request, id):
    try:
        snps = somatic_snp.objects.filter(study_name=id)
        file = CSVFile.objects.get(id=id)
    except germline_snp.DoesNotExist:
        raise Http404('This item does not exist')
    return render(request, 'oncotator/somatic_snpdetail.html', {'snps': snps, 'file': file})

# Germline SNP detail page, displayed when clicked on results

def germline_snp_detail(request, id):
    try:
        snps = germline_snp.objects.filter(study_name=id)
        file = CSVFile.objects.get(id=id)
    except germline_snp.DoesNotExist:
        raise Http404('This item does not exist')
    return render(request, 'oncotator/germline_snpdetail.html', {'snps': snps, 'file': file})


def SomaticReportDraft(request, id):
    try:
        snps = somatic_snp.objects.filter(study_name=id)
    except somatic_snp.DoesNotExist:
        raise Http404('This item does not exist')

    document = Document()
    document.add_heading('Somatic Mutation results')
    document.add_heading('Study Name: %s' % snps[0].study_name, level=2)
    now = datetime.datetime.now()
    document.add_paragraph()
    document.add_paragraph("%s" % now.strftime('%B %d, %Y'))

    document.add_paragraph('Dear %s,'%request.user.first_name)
    document.add_paragraph()
    document.add_paragraph('We have found following 10 interesting SNPs from your query.')
    document.add_paragraph('Please feel free to contact me for any additional information.')
    document.add_paragraph('I look forward to assisting you in this project.')

    document.add_paragraph()
    document.add_paragraph('Best regards,')
    document.add_paragraph('Team Snippy')
    document.add_paragraph()


    table = document.add_table(rows=1, cols=7)
    table.style = 'LightList-Accent6'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Gene'
    hdr_cells[1].text = 'Chromosome'
    hdr_cells[2].text = 'Start'
    hdr_cells[3].text = 'Reference Allele'
    hdr_cells[4].text = 'Alt. Allele'
    hdr_cells[5].text = 'Varient Class'
    hdr_cells[6].text = 'Varient Type'
    for snp in snps[0:10]:
        if snp.ensemblID:
            row_cells = table.add_row().cells
            row_cells[0].text = str(snp.gene)
            row_cells[1].text = str(snp.chromosome)
            row_cells[2].text = str(snp.start)
            row_cells[3].text = str(snp.ref_allele)
            row_cells[4].text = str(snp.alt_allele)
            row_cells[5].text = str(snp.variant_class)
            row_cells[6].text = str(snp.variant_type)

    for snp in snps[0:10]:
        if snp.ensemblID:
            EnsemblID = snp.ensemblID
            server = "http://rest.ensembl.org"
            ext = "/lookup/id/" + EnsemblID + "?expand=1"

            response = requests.get(server + ext, headers={"Content-Type": "application/json"})

            if not response.ok:
                response.raise_for_status()
                sys.exit()

            document.add_paragraph()
            info = json.loads(response.text)
            document.add_heading('Gene: %s' % info["display_name"], level=2)
            document.add_paragraph('Ensembl ID: %s' % info["id"])
            document.add_paragraph(
                'This gene is from %s species and %s assembly.' % (info["species"], info["assembly_name"]))
            document.add_paragraph('It is known as %s and involved in %s.' % (info["description"], info["biotype"]))

    document.save('media/SOMATIC_REPORT_DRAFT.docx')

    try:
        subprocess.check_call(
            ['unoconv', '-f', 'pdf', '-o',
             'media/SOMATIC_REPORT_DRAFT.pdf', '-d', 'document',
             'media/SOMATIC_REPORT_DRAFT.docx'])
    except subprocess.CalledProcessError as e:
        print('CalledProcessError', e)

    return render(request, 'oncotator/somatic_report.html')


def GermlineReportDraft(request, id):
    try:
        snps = germline_snp.objects.filter(study_name=id)
    except germline_snp.DoesNotExist:
        raise Http404('This item does not exist')

    document = Document()
    document.add_heading('Germline Mutation results')
    document.add_heading('Study Name: %s' % snps[0].study_name, level=2)
    now = datetime.datetime.now()
    document.add_paragraph()
    document.add_paragraph("%s" % now.strftime('%B %d, %Y'))

    document.add_paragraph('Dear %s,' % request.user.first_name)
    document.add_paragraph()
    document.add_paragraph('We have found following 10 interesting SNPs from your query.')
    document.add_paragraph('Please feel free to contact me for any additional information.')
    document.add_paragraph('I look forward to assisting you in this project.')

    document.add_paragraph()
    document.add_paragraph('Best regards,')
    document.add_paragraph('Team Snippy')
    document.add_paragraph()

    table = document.add_table(rows=1, cols=7)
    table.style = 'LightList-Accent6'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Gene'
    hdr_cells[1].text = 'Chromosome'
    hdr_cells[2].text = 'Start'
    hdr_cells[3].text = 'Reference Allele'
    hdr_cells[4].text = 'Alt. Allele'
    hdr_cells[5].text = 'Varient Class'
    hdr_cells[6].text = 'Varient Type'
    for snp in snps[0:10]:
        if snp.ensemblID:
            row_cells = table.add_row().cells
            row_cells[0].text = str(snp.gene)
            row_cells[1].text = str(snp.chromosome)
            row_cells[2].text = str(snp.start)
            row_cells[3].text = str(snp.ref_allele)
            row_cells[4].text = str(snp.alt_allele)
            row_cells[5].text = str(snp.variant_class)
            row_cells[6].text = str(snp.variant_type)

    for snp in snps[0:10]:
        if snp.ensemblID:
            EnsemblID = snp.ensemblID
            server = "http://rest.ensembl.org"
            ext = "/lookup/id/" + EnsemblID + "?expand=1"

            response = requests.get(server + ext, headers={"Content-Type": "application/json"})

            if not response.ok:
                response.raise_for_status()
                sys.exit()

            document.add_paragraph()
            info = json.loads(response.text)
            document.add_heading('Gene: %s' % info["display_name"], level=2)
            document.add_paragraph('Ensembl ID: %s' % info["id"])
            document.add_paragraph(
                'This gene is from %s species and %s assembly.' % (info["species"], info["assembly_name"]))
            document.add_paragraph('It is known as %s and involved in %s.' % (info["description"], info["biotype"]))

    document.save('media/GERMLINE_REPORT_DRAFT.docx')

    try:
        subprocess.check_call(
            ['unoconv', '-f', 'pdf', '-o', 'media/GERMLINE_REPORT_DRAFT.pdf', '-d', 'document', 'media/GERMLINE_REPORT_DRAFT.docx'])
    except subprocess.CalledProcessError as e:
        print('CalledProcessError', e)

    #unoconv('media/GERMLINE_REPORT_DRAFT.docx', format='pdf', output='media/GERMLINE_REPORT_DRAFT.pdf')

    return render(request, 'oncotator/germline_report.html')


def germline_gene_network(request, id):
    try:
        snps = germline_snp.objects.filter(study_name=id)
    except germline_snp.DoesNotExist:
        raise Http404('This item does not exist')

    temp_query = 'http://string-db.org/api/image/networkList?identifiers='

    for snp in snps:
        if snp.ensemblID:
            EnsemblID = snp.ensemblID
            temp_query = temp_query + EnsemblID + '%0D'

    query = temp_query + '&network_flavor=confidence&species=9606&required_score=600&limit=5'

    return render(request, 'oncotator/germline_gene_network.html', {'query': query})


def somatic_gene_network(request, id):
    try:
        snps = somatic_snp.objects.filter(study_name=id)
    except somatic_snp.DoesNotExist:
        raise Http404('This item does not exist')

    temp_query = 'http://string-db.org/api/image/networkList?identifiers='

    for snp in snps:
        if snp.ensemblID:
            EnsemblID = snp.ensemblID
            temp_query = temp_query + EnsemblID + '%0D'

    query = temp_query + '&network_flavor=confidence&species=9606&required_score=600&limit=5'
    return render(request, 'oncotator/somatic_gene_network.html', {'query': query})